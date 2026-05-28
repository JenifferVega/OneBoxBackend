"""
Procesa documentos adjuntos (PDF, DOCX, TXT, imágenes) para extraer texto
y, opcionalmente, generar metadatos del proyecto con IA.

Usado tanto desde la web (subida directa) como desde WhatsApp (Twilio Media).
"""
import os
import re
import io
import json
import base64
import uuid
import mimetypes
from datetime import datetime
from typing import Optional, Tuple

import boto3
import requests

from agent.llm import call_llm, extract_json_from_response


# =============================================================================
# Configuración
# =============================================================================

S3_ATTACHMENTS_BUCKET = os.environ.get(
    "S3_ATTACHMENTS_BUCKET", "onebox-attachments-191027238118"
)
AWS_REGION = os.environ.get("AWS_REGION", "us-east-1")

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_FOR_ANALYSIS = 50_000  # caracteres a enviar al LLM para análisis

ALLOWED_EXTENSIONS = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'doc': 'application/msword',
    'txt': 'text/plain',
    'md': 'text/markdown',
    'png': 'image/png',
    'jpg': 'image/jpeg',
    'jpeg': 'image/jpeg',
    'webp': 'image/webp',
}

_s3_client = None


def get_s3_client():
    global _s3_client
    if _s3_client is None:
        _s3_client = boto3.client("s3", region_name=AWS_REGION)
    return _s3_client


# =============================================================================
# Validación
# =============================================================================

def detect_extension(file_name: str, content_type: str = "") -> str:
    """Detecta la extensión normalizada del archivo."""
    if file_name and '.' in file_name:
        ext = file_name.rsplit('.', 1)[-1].lower().strip()
        if ext in ALLOWED_EXTENSIONS:
            return ext
    # Fallback por content-type
    ct = (content_type or '').lower()
    for ext, mime in ALLOWED_EXTENSIONS.items():
        if mime == ct:
            return ext
    if 'pdf' in ct:
        return 'pdf'
    if 'wordprocessingml' in ct or 'docx' in ct:
        return 'docx'
    if 'msword' in ct:
        return 'doc'
    if ct.startswith('image/'):
        if 'png' in ct: return 'png'
        if 'jpeg' in ct or 'jpg' in ct: return 'jpg'
        if 'webp' in ct: return 'webp'
    if ct.startswith('text/'):
        return 'txt'
    return ''


def validate_file(file_bytes: bytes, file_name: str, content_type: str = "") -> Tuple[bool, str, str]:
    """Retorna (valid, ext, error_msg)."""
    if not file_bytes:
        return False, '', 'El archivo está vacío.'
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        return False, '', f'El archivo supera el límite de {MAX_FILE_SIZE_BYTES // (1024*1024)} MB.'
    ext = detect_extension(file_name, content_type)
    if not ext:
        return False, '', f'Formato no soportado. Permitidos: {", ".join(ALLOWED_EXTENSIONS.keys())}.'
    return True, ext, ''


# =============================================================================
# Extracción de texto
# =============================================================================

def extract_text_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        chunks = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or '')
            except Exception:
                continue
        return '\n\n'.join(chunks).strip()
    except Exception as e:
        print(f"[document_parser] PDF error: {e}")
        return ''


def extract_text_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        # Párrafos
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        # Tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(c.text.strip() for c in row.cells if c.text and c.text.strip())
                if row_text:
                    parts.append(row_text)
        return '\n'.join(parts).strip()
    except Exception as e:
        print(f"[document_parser] DOCX error: {e}")
        return ''


def extract_text_txt(file_bytes: bytes) -> str:
    for encoding in ('utf-8', 'utf-16', 'latin-1', 'cp1252'):
        try:
            return file_bytes.decode(encoding).strip()
        except (UnicodeDecodeError, LookupError):
            continue
    return ''


def extract_text_image(file_bytes: bytes, ext: str) -> str:
    """Usa Gemini Vision para extraer texto de una imagen."""
    if not GEMINI_API_KEY:
        return ''
    try:
        b64 = base64.b64encode(file_bytes).decode('utf-8')
        mime = ALLOWED_EXTENSIONS.get(ext, 'image/png')
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
        payload = {
            "contents": [{
                "role": "user",
                "parts": [
                    {"text": "Extrae todo el texto visible en esta imagen y descríbela brevemente. "
                             "Si es un acta, contrato, brief o documento de proyecto, identifica nombre, "
                             "fechas, participantes y objetivos. Responde en español."},
                    {"inline_data": {"mime_type": mime, "data": b64}}
                ]
            }],
            "generationConfig": {"temperature": 0.1, "maxOutputTokens": 4096}
        }
        resp = requests.post(url, json=payload, timeout=60)
        if resp.status_code != 200:
            print(f"[document_parser] Gemini Vision error {resp.status_code}: {resp.text[:200]}")
            return ''
        data = resp.json()
        candidates = data.get("candidates", [])
        if not candidates:
            return ''
        parts = candidates[0].get("content", {}).get("parts", [])
        return '\n'.join(p.get("text", "") for p in parts if "text" in p).strip()
    except Exception as e:
        print(f"[document_parser] image vision error: {e}")
        return ''


def extract_text(file_bytes: bytes, ext: str) -> str:
    """Despacha al extractor según extensión."""
    if ext == 'pdf':
        return extract_text_pdf(file_bytes)
    if ext in ('docx', 'doc'):
        return extract_text_docx(file_bytes)
    if ext in ('txt', 'md'):
        return extract_text_txt(file_bytes)
    if ext in ('png', 'jpg', 'jpeg', 'webp'):
        return extract_text_image(file_bytes, ext)
    return ''


# =============================================================================
# Análisis IA: extraer metadatos del proyecto del texto
# =============================================================================

PROJECT_TYPES = ['Desarrollo Web', 'Infraestructura', 'Diseño', 'Marketing',
                 'Ecommerce', 'Consultoría', 'Soporte', 'RRHH', 'Otro']


def analyze_document_for_project(text: str, fallback_name: str = "") -> dict:
    """
    Pide a la IA que infiera metadatos del proyecto a partir del texto del documento.
    Retorna {name, type, description, extractedNotes}.
    """
    if not text or len(text.strip()) < 30:
        return {
            "name": fallback_name or "Proyecto sin título",
            "type": "Otro",
            "description": text.strip()[:500] if text else "",
            "extractedNotes": ""
        }

    # Limitar el texto enviado al LLM
    truncated = text[:MAX_TEXT_FOR_ANALYSIS]
    types_str = ', '.join(PROJECT_TYPES)

    prompt = f"""Analiza el siguiente documento o conversación y propone metadatos para crear un proyecto. Tu objetivo es generar una descripción FÁCTICA del PROYECTO en sí, NO de la conversación entre los autores.

DOCUMENTO O CONVERSACIÓN:
{truncated}

REGLAS CRÍTICAS PARA "description":
1. La descripción debe ser sobre EL PROYECTO (qué es, qué tiene, qué stack, qué entregables), NO sobre la conversación entre las personas.
   - ❌ INCORRECTO: "Kevin solicitó a Mateo un proyecto..." / "Santi pasó las credenciales a Belen..."
   - ✅ CORRECTO: "GhostLink, aplicación de mensajería..." / "Actualización del sitio WordPress les-sp.org..."
2. NO narres "X dijo a Y", "X propuso a Y", "X solicitó". Describe el proyecto como un brief técnico/ejecutivo.
3. PROHIBIDO usar lenguaje genérico ("una iniciativa para optimizar", "una serie de mejoras", "un desafío importante"). Eso NO informa.
4. OBLIGATORIO mencionar hechos concretos del proyecto:
   - Nombre del producto/proyecto (ej: "GhostLink", "les-sp.org")
   - Funcionalidades específicas con detalles
   - Stack tecnológico exacto si aparece (React, Node.js, MongoDB, etc.)
   - URLs / dominios / plataformas mencionadas
   - Tareas concretas con detalles
   - Métricas exactas (horas, €, fechas)
   - Decisiones puntuales sobre el producto
5. Las personas que solo participan en la conversación (autores del chat) NO van en la descripción.
   - Excepción: si una persona es el cliente final, responsable o rol clave del proyecto, sí incluirla.
6. 5-8 oraciones, densas en información concreta sobre el PROYECTO.

EJEMPLO DE BUENA description (referencia, NO copiar):
"GhostLink: aplicación de mensajería con temática cyberpunk inspirada visualmente en WhatsApp pero orientada a 'agentes secretos'. Funcionalidades clave: chats en tiempo real, estados, grupos, audios, videollamadas fake y mensajes autodestruibles. Interfaz: sidebar de chats, mensajes alineados derecha/izquierda, checks de enviado/leído, barra de escritura con emojis y modo oscuro obligatorio en negro con verde neón. Login y perfiles de usuario como 'agentes' con alias, foto y estado. Stack: React + Next.js + Tailwind en frontend; Node.js + Express + Socket.io en backend; MongoDB como base de datos. El producto se concibe como una demo."

EJEMPLO DE MALA description (NO HAGAS ESTO):
"Kevin solicitó a Mateo un proyecto de aplicación de mensajería. Mateo propuso GhostLink. Las funcionalidades incluyen chats..." (esto narra la conversación, no describe el proyecto)

REGLAS PARA "name":
- Usa el nombre del producto/proyecto mencionado en el texto (ej: "GhostLink", "LES España y Portugal").
- NO uses los nombres de las personas que conversaron como nombre del proyecto.
- Máximo 60 caracteres.

REGLAS PARA "type":
- Elige una categoría exacta: {types_str}
- Que refleje el alcance REAL (no solo el solicitado).

REGLAS PARA "extractedNotes":
- Caracterización breve del proyecto (máx 200 chars).
- Ejemplo: "App móvil con stack MERN y temática cyberpunk para demo de producto".

REGLAS PARA "detected_participants":
- Lista las personas mencionadas en el texto (autores del chat, clientes, miembros del equipo).
- Cada uno con: name (el nombre exacto), role_inferred (rol inferido del contexto, ej: "Cliente", "Desarrollador", "Diseñador", "PM", "Responsable técnico").
- Si no hay personas claras, devuelve lista vacía.
- Máximo 8 personas.

NO INVENTES información que no esté en el texto.

RESPONDE SOLO JSON, sin texto extra ni bloques de código:
{{
  "name": "...",
  "type": "...",
  "description": "...",
  "extractedNotes": "...",
  "detected_participants": [
    {{"name": "Kevin", "role_inferred": "Solicitante / Cliente"}},
    {{"name": "Mateo", "role_inferred": "Desarrollador"}}
  ]
}}"""

    try:
        response = call_llm(
            system_prompt="Eres un analista que extrae datos CONCRETOS y FÁCTICOS de un texto: nombres propios, URLs, números, tareas concretas, decisiones específicas. Prohibido el lenguaje genérico. Devuelves SIEMPRE JSON válido en español sin bloques de código markdown.",
            user_message=prompt,
            temperature=0.15,
            max_tokens=2048
        )
        analysis = extract_json_from_response(response)
        if not analysis:
            cleaned = response.strip()
            if cleaned.startswith('```'):
                cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
                cleaned = re.sub(r'\s*```$', '', cleaned)
            try:
                analysis = json.loads(cleaned)
            except json.JSONDecodeError:
                last_brace = cleaned.rfind('}')
                if last_brace > 0:
                    analysis = json.loads(cleaned[:last_brace + 1])
                else:
                    raise

        # Validaciones / valores por defecto
        name = str(analysis.get('name', '')).strip()[:80] or fallback_name or "Proyecto sin título"
        ptype = str(analysis.get('type', 'Otro')).strip()
        if ptype not in PROJECT_TYPES:
            ptype = 'Otro'
        description = str(analysis.get('description', '')).strip()
        if not description:
            description = text.strip()[:500]
        notes = str(analysis.get('extractedNotes', '')).strip()[:300]

        # Procesar detected_participants
        raw_participants = analysis.get('detected_participants', []) or []
        detected_participants = []
        seen_names = set()
        for p in raw_participants[:8]:
            if not isinstance(p, dict):
                continue
            pname = str(p.get('name', '')).strip()[:80]
            prole = str(p.get('role_inferred', '')).strip()[:80]
            if pname and pname.lower() not in seen_names:
                seen_names.add(pname.lower())
                detected_participants.append({
                    'name': pname,
                    'role_inferred': prole or 'Participante'
                })

        return {
            "name": name,
            "type": ptype,
            "description": description,
            "extractedNotes": notes,
            "detected_participants": detected_participants
        }
    except Exception as e:
        print(f"[document_parser] analyze error: {e}")
        return {
            "name": fallback_name or "Proyecto desde documento",
            "type": "Otro",
            "description": text.strip()[:500],
            "extractedNotes": "",
            "detected_participants": []
        }


# =============================================================================
# Almacenamiento en S3
# =============================================================================

def upload_to_s3(file_bytes: bytes, project_id: str, file_name: str, content_type: str = "application/octet-stream") -> str:
    """Sube el archivo a S3 y retorna la s3 key."""
    safe_name = re.sub(r'[^\w\s\-\.]', '_', file_name)[:200]
    key = f"projects/{project_id}/{datetime.utcnow().strftime('%Y%m%d')}/{uuid.uuid4().hex[:8]}_{safe_name}"
    client = get_s3_client()
    client.put_object(
        Bucket=S3_ATTACHMENTS_BUCKET,
        Key=key,
        Body=file_bytes,
        ContentType=content_type or 'application/octet-stream',
        ServerSideEncryption='AES256'
    )
    print(f"[document_parser] S3 upload OK: {key} ({len(file_bytes)} bytes)")
    return key


def generate_download_url(s3_key: str, file_name: str = "", expires_in: int = 600) -> str:
    """Genera una URL presignada para descargar un archivo (válida por X segundos)."""
    client = get_s3_client()
    params = {
        'Bucket': S3_ATTACHMENTS_BUCKET,
        'Key': s3_key,
    }
    if file_name:
        params['ResponseContentDisposition'] = f'attachment; filename="{file_name}"'
    return client.generate_presigned_url('get_object', Params=params, ExpiresIn=expires_in)


def delete_from_s3(s3_key: str) -> bool:
    try:
        client = get_s3_client()
        client.delete_object(Bucket=S3_ATTACHMENTS_BUCKET, Key=s3_key)
        return True
    except Exception as e:
        print(f"[document_parser] S3 delete error: {e}")
        return False


# =============================================================================
# Twilio: descargar archivo de un MediaUrl
# =============================================================================

def download_from_twilio(media_url: str, account_sid: str, auth_token: str) -> Tuple[bytes, str, str]:
    """Descarga un archivo desde una MediaUrl de Twilio.
    Retorna (file_bytes, content_type, filename)."""
    try:
        resp = requests.get(media_url, auth=(account_sid, auth_token), timeout=60, allow_redirects=True)
        if resp.status_code != 200:
            print(f"[document_parser] Twilio download error {resp.status_code}: {resp.text[:200]}")
            return b'', '', ''
        ct = resp.headers.get('Content-Type', 'application/octet-stream').split(';')[0].strip()
        # Inferir nombre
        ext = mimetypes.guess_extension(ct) or ''
        if ext.startswith('.'):
            ext = ext[1:]
        if not ext:
            ext = detect_extension('', ct)
        fname = f"whatsapp_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.{ext or 'bin'}"
        return resp.content, ct, fname
    except Exception as e:
        print(f"[document_parser] Twilio download exc: {e}")
        return b'', '', ''
